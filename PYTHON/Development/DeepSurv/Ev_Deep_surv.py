import pandas as pd
import numpy as np
import argparse
import matplotlib.pyplot as plt 
from sklearn.utils import resample
import os
import numpy as np
import dask.distributed
import pycox
import matplotlib.pyplot as plt
import torch
from pycox.models import CoxPH
from sklearn.metrics import classification_report, confusion_matrix, roc_curve,auc
from docx import Document
import io
import sys
from sklearn.utils import resample
import sksurv.metrics
from docx.shared import Inches
####################### FUNCTION FOR CALIBRATION PLOT WITH CI #############
def bootstrap_calibration(predicted_probs, y_test_event,n_bootstrap = 100, n_bins =10, prob_max = 0.03):
    
    bins = np.linspace(0, prob_max, n_bins+1 )
    mean_pred_boot = []
    mean_obs_boot = []
    
    for s in range(n_bootstrap):
        predicted_probs_resample, y_test_event_resample = resample (predicted_probs, y_test_event)
        mean_predicted = []
        mean_observed = []
        bin_indices = np.digitize(predicted_probs_resample, bins) - 1
        for i in range(n_bins):

            bin_mask = bin_indices == i
            if np.sum(bin_mask) > 0:
                mean_predicted.append(np.mean(predicted_probs_resample[bin_mask]))
                observed_in_bin = np.mean(y_test_event_resample[bin_mask])
                mean_observed.append(observed_in_bin)
            else:
                mean_predicted.append(np.nan)
                mean_observed.append(np.nan)
        

        mean_predicted = np.array(mean_predicted)
        mean_observed = np.array(mean_observed)
        
        mean_pred_boot.append (mean_predicted)
        mean_obs_boot.append(mean_observed)
    
    mean_pred_boot_low =  np.percentile(mean_pred_boot, 2.5, axis=0)
    mean_pred_boot_upp =  np.percentile(mean_pred_boot, 97.5, axis=0)

    mean_obs_boot_low =  np.percentile(mean_obs_boot, 2.5, axis=0)
    mean_obs_boot_upp =  np.percentile(mean_obs_boot, 97.5, axis=0)

    df_cal = pd.DataFrame( {   "Cal_pred_low" : mean_pred_boot_low,
                                "Cal_pred_upp" : mean_pred_boot_upp,
                                "Cal_obs_low"  : mean_obs_boot_low,
                                "Cal_obs_upp"  : mean_obs_boot_upp })
    return df_cal

def bootstrap_metrics (predictions, true_event , n_bootstrap = 100):

    boot_sens = []
    boot_spec = []
    boot_ppv = []
    boot_npv = []
    
    for i in range(n_bootstrap):
        predictions_resampled , true_event_resampled = resample( predictions, true_event)
        conf_matrix = confusion_matrix (true_event_resampled, predictions_resampled)

        tp = conf_matrix[1, 1]
        fn = conf_matrix[1, 0]
        tn = conf_matrix[0, 0]
        fp = conf_matrix[0, 1]

        boot_sens.append( tp / (tp + fn))
        boot_spec.append ( tn / (tn + fp))
        boot_ppv.append ( tp / ( tp+ fp))
        boot_npv.append (tn / (tn + fn))

    boot_sens_low =  np.percentile(boot_sens, 2.5, axis=0)
    boot_sens_upp =  np.percentile(boot_sens, 97.5, axis=0)

    boot_spec_low =  np.percentile(boot_spec, 2.5, axis=0)
    boot_spec_upp =  np.percentile(boot_spec, 97.5, axis=0)

    boot_ppv_low =  np.percentile(boot_ppv, 2.5, axis=0)
    boot_ppv_upp =  np.percentile(boot_ppv, 97.5, axis=0)

    boot_npv_low =  np.percentile(boot_npv, 2.5, axis=0)
    boot_npv_upp =  np.percentile(boot_npv, 97.5, axis=0)

    df_metric = pd.DataFrame( { "Sens_low" : boot_sens_low,
                                "Sens_upp" : boot_sens_upp,
                                "Spec_low" : boot_spec_low,
                                "Spec_upp" : boot_spec_upp,
                                "PPV_low"  : boot_ppv_low,
                                "PPV_upp"  : boot_ppv_upp,
                                "NPV_low"  : boot_npv_low,
                                "NPV_upp"  : boot_npv_upp }, index = [0])
    return df_metric
####################### FUNCTION FOR CIS FOR AUC #############
def bootstrap_auc(predicted_prob,y_test_true_event, time ,n_bootstrap=200):
    boot_auc = []
    for i in range(n_bootstrap):
        predicted_probs_resample, y_test_true_resampled = resample(predicted_prob, 
        y_test_true_event)
        fpr, tpr, threshold = roc_curve( y_test_true_resampled, predicted_probs_resample)
        boot_auc.append (auc(fpr,tpr))
    low_bound = np.percentile(boot_auc, 2.5, axis=0)
    upper_bound = np.percentile(boot_auc, 97.5, axis = 0)
    return low_bound, upper_bound
####################### FUNCTION FOR CREATING THE REPORT #############
def create_classification_report( true_col, pred_col):

    y_true = true_col
    y_pred = pred_col

    print("Classifier Report")
    print(classification_report(y_true, y_pred))

    print("\nConfusion Matrix")
    conf_matrix = confusion_matrix(y_true, y_pred)
    print(conf_matrix)

    # Sensitivity (Recall for class 1)
    tp = conf_matrix[1, 1]  # True positives (class 1)
    fn = conf_matrix[1, 0]  # False negatives (class 1)
    sensitivity = tp / (tp + fn)  # Sensitivity formula
    print(f"\nSensitivity: {sensitivity:.2f}")

    # Specificity (Recall for class 0)
    tn = conf_matrix[0, 0]  # True negatives (class 0)
    fp = conf_matrix[0, 1]  # False positives (class 0)
    specificity = tn / (tn + fp)  # Specificity formula
    print(f"Specificity: {specificity:.2f}")

def add_content_to_doc(col_true, col_pred, heading,optimal_cutoff):
    captured_output = io.StringIO()
    sys.stdout = captured_output 
    create_classification_report( col_true,col_pred )
    print(f"Optimal Cut-off: {optimal_cutoff:.2f}")
    confusion_matrix(col_true, col_pred)
    print (f"Calibration-in-the-large index= {np.mean(col_pred) - np.mean(col_true):.4f}")
    sys.stdout = sys.__stdout__
    printed_text = captured_output.getvalue()
    doc.add_heading(heading, 0)
    doc.add_paragraph(printed_text)

# Calculate the optimal cut-off
def sensivity_specifity_cutoff(y_test, y_pred_nn):
    fpr, tpr, thresholds = roc_curve(y_test, y_pred_nn)
    idx = np.argmax(tpr - fpr)
    return thresholds[idx]

# Function to add classification report as a table
def add_classification_report_to_doc(doc, report, year):
    doc.add_paragraph(f"Classification Report for Year {year}:", style="Heading 3")
    table = doc.add_table(rows=1, cols=5)
    table.style = "Light Grid Accent 1"

    # Add header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Class"
    hdr_cells[1].text = "Precision"
    hdr_cells[2].text = "Recall"
    hdr_cells[3].text = "F1-Score"
    hdr_cells[4].text = "Support"

    # Add data rows
    for key, value in report.items():
        if isinstance(value, dict):
            row_cells = table.add_row().cells
            row_cells[0].text = key
            row_cells[1].text = f"{value['precision']:.2f}"
            row_cells[2].text = f"{value['recall']:.2f}"
            row_cells[3].text = f"{value['f1-score']:.2f}"
            row_cells[4].text = f"{value['support']}"

    # Add overall metrics (accuracy, macro avg, weighted avg)
    for key in ["accuracy", "macro avg", "weighted avg"]:
        if key in report:
            row_cells = table.add_row().cells
            row_cells[0].text = key
            row_cells[1].text = "-"
            row_cells[2].text = "-"
            row_cells[3].text = f"{report[key]['f1-score']:.2f}" if isinstance(report[key], dict) else "-"
            row_cells[4].text = f"{report[key]:.2f}" if key == "accuracy" else "-"
    

    # Function to add confusion matrix as a table
def add_confusion_matrix_to_doc(doc, conf_matrix, year):
    doc.add_paragraph(f"Confusion Matrix for Year {year}:", style="Heading 3")
    table = doc.add_table(rows=conf_matrix.shape[0] + 1, cols=conf_matrix.shape[1] + 1)
    table.style = "Light Grid Accent 1"

    # Add header row
    header_cells = table.rows[0].cells
    header_cells[0].text = ""
    for i in range(conf_matrix.shape[1]):
        header_cells[i + 1].text = f"Predicted {i}"

    # Add data rows
    for i in range(conf_matrix.shape[0]):
        row_cells = table.rows[i + 1].cells
        row_cells[0].text = f"Actual {i}"
        for j in range(conf_matrix.shape[1]):
            row_cells[j + 1].text = str(conf_matrix[i, j])

# Main function
def main(MODEL_NAME, data_path,data_path_sub, model_path, output_dir, max_years):

    # READ DATA
    df = pd.read_parquet(data_path, engine="pyarrow")
    df_train = df[df['Sub_Cohort'] == 'Discovery']
    del (df)
    df = pd.read_parquet(data_path_sub, engine="pyarrow")
    df_val =  df[df['Sub_Cohort'] == 'Validation']
    del(df)
    df_train = df_train.drop(columns=['Sub_Cohort'])
    df_val = df_val.drop(columns=['Sub_Cohort'])

    # PROCESS DATA
    get_target = lambda df: (df['Inc_Fx_Inc_fx_hip_days_to_cens'].values, df['Inc_Fx_Inc_fx_hip_bin_cens'].values)
    get_feat = lambda df: (df.drop(['Inc_Fx_Inc_fx_hip_days_to_cens','Inc_Fx_Inc_fx_hip_bin_cens'],axis=1).values)
    df_train = df_train.dropna()
    df_val = df_val.dropna()
    y_train = get_target(df_train)
    y_val = get_target(df_val)
    x_train = get_feat (df_train)
    x_val = get_feat ( df_val)
    x_train = x_train.astype("float32")
    x_val = x_val.astype("float32")

    # Model
    net=torch.load(model_path)
    model = CoxPH(net)
    model.net.eval()  

    # Evaluation
    doc = Document()
    doc.add_heading("DeepSurv model Evaluation", level=1)
    _= model.compute_baseline_hazards(x_val, y_val)
    surv = model.predict_surv_df(x_val)
    _= model.compute_baseline_hazards(x_train, y_train)
    surv_train = model.predict_surv_df(x_train)
    doc = Document()
    surv = surv.dropna(axis=1)
    y_true = df_val    [["Inc_Fx_Inc_fx_hip_bin_cens"]]
    df_pred = y_true
    df_pred_train = df_train    [["Inc_Fx_Inc_fx_hip_bin_cens"]]

    sens = []
    spec = []
    PPV = []
    NPV = []
    auc_scores = []
    auc_low = []
    auc_upp = []

    for N in range(1,max_years):
        days = 365*N
        predicted_probs = 1 - surv.iloc[days].values
        predicted_probs_train = 1 - surv_train.iloc[days].values
        df_pred["Event_obs_year"+str(N)] = [1 if val <365*N else 0 for val in df_val["Inc_Fx_Inc_fx_hip_days_to_cens"]]
        df_pred.loc[df_val["Inc_Fx_Inc_fx_hip_bin_cens"]==0,"Event_obs_year"+str(N)]=0
        y_test_event =  df_pred["Event_obs_year"+str(N)]
        df_pred_train["Event_obs_year"+str(N)] = [1 if val <365*N else 0 for val in df_train["Inc_Fx_Inc_fx_hip_days_to_cens"]]
        df_pred_train.loc[df_train["Inc_Fx_Inc_fx_hip_bin_cens"]==0,"Event_obs_year"+str(N)]=0
        y_test_event_train =  df_pred_train["Event_obs_year"+str(N)]
        optimal_cutoff = sensivity_specifity_cutoff(y_test_event_train, predicted_probs_train)

        print(f"Years: {N}")
        doc.add_heading(f"Year {N} Results", level=2)
        doc.add_paragraph(f"Optimal Cut-off: {optimal_cutoff:.4f}")

        fpr, tpr, threshold = roc_curve( y_test_event, predicted_probs)
        roc_auc = auc(fpr,tpr)
        auc_scores.append(roc_auc)
        low, upp = bootstrap_auc(predicted_probs,y_test_event, days ,n_bootstrap=200)
        auc_low.append(low)
        auc_upp.append(upp)

        doc.add_paragraph(f"AUC: {roc_auc:.4f} ({low:.4f}, {upp:.4f})")
        plt.plot(fpr, tpr, lw=2, label=f'ROC curve - Year {N} (AUC = {roc_auc:.2f})')
        
        # Apply the optimal cut-off
        binary_predictions = (predicted_probs >= optimal_cutoff).astype(int)
        df_met = bootstrap_metrics (binary_predictions, y_test_event)

        doc.add_paragraph(f"Calibration-in-the-large index: {np.mean(binary_predictions) - np.mean(y_test_event):.4f}")
        report = classification_report(y_test_event, binary_predictions, 
        target_names=["No Event", "Event"], output_dict=True)
        add_classification_report_to_doc(doc, report, N)

        # Confusion matrix
        conf_matrix = confusion_matrix(y_test_event, binary_predictions)
        add_confusion_matrix_to_doc(doc, conf_matrix, N)

        tp = conf_matrix[1, 1]  # True positives (class 1)
        fn = conf_matrix[1, 0]  # False negatives (class 1)
        sens.append( tp / (tp + fn))  # Sensitivity formula
        sensitiv = tp / (tp + fn)
        doc.add_paragraph(f"Sensitivity: {sensitiv:.4f} ({df_met.iloc[0,0]:.4f}, {df_met.iloc[0,1]:.4f})")
    
        tn = conf_matrix[0, 0]  # True negatives (class 0)
        fp = conf_matrix[0, 1]  # False positives (class 0)
        spec.append( tn / (tn + fp))  # Specificity formula
        specific = tn / (tn + fp)
        doc.add_paragraph(f"Specificity: {specific:.4f} ({df_met.iloc[0,2]:.4f}, {df_met.iloc[0,3]:.4f})") 

        # PPV
        ppv_value  =  tp / (tp + fp)
        PPV.append ( ppv_value ) 
        doc.add_paragraph(f"Positive Predictive Value: {ppv_value:.4f} ({df_met.iloc[0,4]:.4f}, {df_met.iloc[0,5]:.4f})")

        # NPV
        npv_value = tn / (tn + fn)
        NPV.append( npv_value )
        doc.add_paragraph(f"Negative Predictive Value: {npv_value:.4f} ({df_met.iloc[0,6]:.4f}, {df_met.iloc[0,7]:.4f})")

    plt.xlabel('False Positive Rate')
    plt.ylabel('True Positive Rate')
    plt.title('ROC Curve for All Years')
    plt.legend(loc="lower right")
    plt.grid(alpha=0.3)
    roc_buffer = io.BytesIO()
    plt.savefig(roc_buffer, format="png")
    plt.savefig(output_dir+ "/ROC_DeepSurv_"+MODEL_NAME+".pdf", format='pdf', dpi=1200)  # High-quality PDF format
    plt.savefig(output_dir+ "/ROC_DeepSurv_"+MODEL_NAME+".eps", format='eps', dpi=1200)
    plt.close()

    doc.add_heading("ROC Curve for All Years", level=1)
    doc.add_picture(roc_buffer, width=Inches(6))
    roc_buffer.close()

    # AUC plot
    plt.figure(figsize=(10, 6))
    plt.plot(range(1, max_years), auc_scores, marker="o", label="AUC")
    plt.plot(range(1, max_years), sens, marker="o", label="Sensitivity")
    plt.plot(range(1, max_years), spec, marker="o", label="Specificity")
    plt.plot(range(1, max_years), PPV, marker="o", label="PPV")
    plt.plot(range(1, max_years), NPV, marker="o", label="NPV")
    plt.xlabel("Years")
    plt.ylabel("Metric Value")
    plt.title("AUC, Sensitivity, Specificity, PPV and NPV Over 10 Years")
    plt.legend()
    plt.grid(alpha=0.3)
    metrics_buffer = io.BytesIO()
    plt.savefig(metrics_buffer, format="png")
    plt.savefig(output_dir + "/AUC_DeepSurv_"+MODEL_NAME+".pdf", format='pdf', dpi=1200)  # High-quality PDF format
    plt.savefig(output_dir + "/AUC_DeepSurv_"+MODEL_NAME+".eps", format='eps', dpi=1200)
    plt.close()

    doc.add_heading("AUC, Sensitivity, Specificity, PPV and NPC Over 10 Years", level=1)
    doc.add_picture(metrics_buffer, width=Inches(6))
    metrics_buffer.close()

    # Calubration plot
    for N in range(1,max_years):
        days = 365*N
        predicted_probs = 1-surv.iloc[days].values
        df_pred["Event_obs_year"+str(N)] = [1 if val <365*N else 0 for val in df_val["Inc_Fx_Inc_fx_hip_days_to_cens"]]
        df_pred.loc[df_val["Inc_Fx_Inc_fx_hip_bin_cens"]==0,"Event_obs_year"+str(N)]=0
        y_test_event =  df_pred["Event_obs_year"+str(N)]
        n_bins = 10
        bins = np.linspace(0, 0.03, n_bins + 1)
        bin_indices = np.digitize(predicted_probs, bins) - 1

        mean_predicted = []
        mean_observed = []

        for i in range(n_bins):
            # Get indices for current bin
            bin_mask = bin_indices == i
            if np.sum(bin_mask) > 0:

                # Mean predicted probability for the bin
                mean_predicted.append(np.mean(predicted_probs[bin_mask]))

                # Observed survival probability (Kaplan-Meier at the chosen time)
                observed_in_bin = np.mean(y_test_event[bin_mask])

                mean_observed.append(observed_in_bin)
            else:
                # Skip empty bins
                mean_predicted.append(np.nan)
                mean_observed.append(np.nan)

        # Drop NaNs for empty bins
        mean_predicted = np.array(mean_predicted)
        mean_observed = np.array(mean_observed)
        valid_mask = ~np.isnan(mean_predicted)
        df_cal = bootstrap_calibration ( predicted_probs, y_test_event)

        plt.figure(figsize=(8, 6))
        plt.plot(mean_predicted[valid_mask], mean_observed[valid_mask],  marker='o')
        plt.plot([0, 0.03], [0, 0.03], 'k--')
        plt.xlabel("Predicted Survival Probability")
        plt.ylabel("Observed Survival Probability")
        plt.fill_between(mean_predicted[valid_mask], df_cal["Cal_obs_low"], df_cal["Cal_obs_upp"], color = "b",alpha=0.2,label="95% CI")
        plt.title(f"Calibration Plot at Year {N}")
        plt.legend()
        plt.grid(alpha=0.3)
        metrics_buffer = io.BytesIO()
        plt.savefig(metrics_buffer, format="png")
        plt.savefig(output_dir+"/Calibration_"+MODEL_NAME+"Year"+str(N)+".pdf", format='pdf', dpi=1200)  # High-quality PDF format
        plt.savefig(output_dir+"/Calibration_"+MODEL_NAME+"Year"+str(N)+".eps", format='eps', dpi=1200)
        plt.close()

        doc.add_heading(f"Calibration curve Year {N}", level=1)
        doc.add_picture(metrics_buffer, width=Inches(6))
        metrics_buffer.close()

    doc.save(output_dir+"/DeepSurv_"+MODEL_NAME+"f.docx")

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Evaluate Deep surv model.")
    parser.add_argument("--MODEL_NAME", type=str, required=True, help="Model names.")
    parser.add_argument("--data_path", type=str, required=True, help="Directory where the data are.")
    parser.add_argument("--data_path_sub", type=str, required=True, help="Directory where the data for the subgroup are.")
    parser.add_argument("--model_path", type=str, required=True, help="Directory where the model is saved.")
    parser.add_argument("--output_dir", type=str, required=True, help="Directory where the output is saved.")
    parser.add_argument("--max_years", type=int, required=True, help="Maximum number of years+1.")
    args = parser.parse_args()

    # Run the main function with the provided arguments
    main(args.MODEL_NAME, args.data_path, args.data_path_sub,args.model_path, args.output_dir,args.max_years)