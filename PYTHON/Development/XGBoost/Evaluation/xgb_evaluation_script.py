import os
import pandas as pd
import numpy as np
import xgboost as xgb
from scipy.stats import norm
import matplotlib.pyplot as plt
from sklearn.metrics import classification_report, confusion_matrix, roc_curve, auc
from docx import Document
from docx.shared import Inches
import io
import argparse
from sklearn.utils import resample

def bootstrap_calibration(predicted_probs, y_test_event,n_bootstrap = 100, n_bins =10, prob_max = 0.03):
    
    bins = np.linspace(0, prob_max, n_bins+1 )
    mean_pred_boot = []
    mean_obs_boot = []
    
    for s in range(n_bootstrap):
        predicted_probs_resample, y_test_event_resample = resample (predicted_probs, y_test_event)
        mean_predicted = []
        mean_observed = []
        bin_indices = np.digitize(predicted_probs_resample, bins) - 1
        for i in range(n_bins):

            bin_mask = bin_indices == i
            if np.sum(bin_mask) > 0:
                mean_predicted.append(np.mean(predicted_probs_resample[bin_mask]))
                observed_in_bin = np.mean(y_test_event_resample[bin_mask])
                mean_observed.append(observed_in_bin)
            else:
                mean_predicted.append(np.nan)
                mean_observed.append(np.nan)
        

        mean_predicted = np.array(mean_predicted)
        mean_observed = np.array(mean_observed)
        

        mean_pred_boot.append (mean_predicted)
        mean_obs_boot.append(mean_observed)
    
    mean_pred_boot_low =  np.percentile(mean_pred_boot, 2.5, axis=0)
    mean_pred_boot_upp =  np.percentile(mean_pred_boot, 97.5, axis=0)

    mean_obs_boot_low =  np.percentile(mean_obs_boot, 2.5, axis=0)
    mean_obs_boot_upp =  np.percentile(mean_obs_boot, 97.5, axis=0)

    df_cal = pd.DataFrame( {   "Cal_pred_low" : mean_pred_boot_low,
                                "Cal_pred_upp" : mean_pred_boot_upp,
                                "Cal_obs_low"  : mean_obs_boot_low,
                                "Cal_obs_upp"  : mean_obs_boot_upp })

    return df_cal

def bootstrap_metrics (predictions, true_event , n_bootstrap = 100):
    boot_sens = []
    boot_spec = []
    boot_ppv = []
    boot_npv = []
    

    for i in range(n_bootstrap):
        predictions_resampled , true_event_resampled = resample( predictions, true_event)
        conf_matrix = confusion_matrix (true_event_resampled, predictions_resampled)

        tp = conf_matrix[1, 1]  # True negatives (class 0)
        fn = conf_matrix[1, 0]
        tn = conf_matrix[0, 0]  # True negatives (class 0)
        fp = conf_matrix[0, 1]

        boot_sens.append( tp / (tp + fn))
        boot_spec.append ( tn / (tn + fp))
        boot_ppv.append ( tp / ( tp+ fp))
        boot_npv.append (tn / (tn + fn))

    boot_sens_low =  np.percentile(boot_sens, 2.5, axis=0)
    boot_sens_upp =  np.percentile(boot_sens, 97.5, axis=0)

    boot_spec_low =  np.percentile(boot_spec, 2.5, axis=0)
    boot_spec_upp =  np.percentile(boot_spec, 97.5, axis=0)

    boot_ppv_low =  np.percentile(boot_ppv, 2.5, axis=0)
    boot_ppv_upp =  np.percentile(boot_ppv, 97.5, axis=0)

    boot_npv_low =  np.percentile(boot_npv, 2.5, axis=0)
    boot_npv_upp =  np.percentile(boot_npv, 97.5, axis=0)

    df_metric = pd.DataFrame( { "Sens_low" : boot_sens_low,
                                "Sens_upp" : boot_sens_upp,
                                "Spec_low" : boot_spec_low,
                                "Spec_upp" : boot_spec_upp,
                                "PPV_low"  : boot_ppv_low,
                                "PPV_upp"  : boot_ppv_upp,
                                "NPV_low"  : boot_npv_low,
                                "NPV_upp"  : boot_npv_upp }, index = [0])

    return df_metric
def bootstrap_auc(predicted_prob,y_test_true_event, time ,n_bootstrap=200):
    boot_auc = []
    for i in range(n_bootstrap):
        predicted_probs_resample, y_test_true_resampled = resample(predicted_prob, 
        y_test_true_event)
        fpr, tpr, threshold = roc_curve( y_test_true_resampled, predicted_probs_resample)
        boot_auc.append (auc(fpr,tpr))
    low_bound = np.percentile(boot_auc, 2.5, axis=0)
    upper_bound = np.percentile(boot_auc, 97.5, axis = 0)
    return low_bound, upper_bound

def main(input_dir, model, output_dir,output_fig):
    dff = pd.read_parquet(input_dir)

    val = dff[dff['Sub_Cohort'] == 'Validation']
    del(dff)
    val = val.drop(columns=['Sub_Cohort'])

    outcome_columns = ["Inc_Fx_Inc_fx_hip_bin_cens", "Inc_Fx_Inc_fx_hip_days_to_cens"]
    predictors = val.columns.difference(outcome_columns).tolist()

    val['label_lower'] = val['Inc_Fx_Inc_fx_hip_days_to_cens']
    val['label_upper'] = np.where(val['Inc_Fx_Inc_fx_hip_bin_cens']==1, val['Inc_Fx_Inc_fx_hip_days_to_cens'], 4200)

    X = val[predictors].values
    y_lower = val['label_lower'].values
    y_upper = val['label_upper'].values

    X_df = pd.DataFrame(X, columns=predictors)
    dval = xgb.DMatrix(X_df, feature_names=predictors)
    dval.set_float_info('label_lower_bound', y_lower)
    dval.set_float_info('label_upper_bound', y_upper)

    bst = xgb.Booster()
    bst.load_model(model)

    # Define survival function
    scale = 1.0
    sf_func = lambda t, pred: norm.sf((np.log(t) - np.log(pred)) / scale)

    # Function to calculate optimal cutoff
    def sensivity_specifity_cutoff(y_test, y_pred_nn):
        fpr, tpr, thresholds = roc_curve(y_test, y_pred_nn)
        idx = np.argmax(tpr - fpr)
        return thresholds[idx]

    # Function to add classification report as a table
    def add_classification_report_to_doc(doc, report, year):
        doc.add_paragraph(f"Classification Report for Year {year}:", style="Heading 3")
        table = doc.add_table(rows=1, cols=5)
        table.style = "Light Grid Accent 1"

        # Add header row
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Class"
        hdr_cells[1].text = "Precision"
        hdr_cells[2].text = "Recall"
        hdr_cells[3].text = "F1-Score"
        hdr_cells[4].text = "Support"

        # Add data rows
        for key, value in report.items():
            if isinstance(value, dict):
                row_cells = table.add_row().cells
                row_cells[0].text = key
                row_cells[1].text = f"{value['precision']:.2f}"
                row_cells[2].text = f"{value['recall']:.2f}"
                row_cells[3].text = f"{value['f1-score']:.2f}"
                row_cells[4].text = f"{value['support']}"

        # Add overall metrics (accuracy, macro avg, weighted avg)
        for key in ["accuracy", "macro avg", "weighted avg"]:
            if key in report:
                row_cells = table.add_row().cells
                row_cells[0].text = key
                row_cells[1].text = "-"
                row_cells[2].text = "-"
                row_cells[3].text = f"{report[key]['f1-score']:.2f}" if isinstance(report[key], dict) else "-"
                row_cells[4].text = f"{report[key]:.2f}" if key == "accuracy" else "-"

    # Function to add confusion matrix as a table
    def add_confusion_matrix_to_doc(doc, conf_matrix, year):
        doc.add_paragraph(f"Confusion Matrix for Year {year}:", style="Heading 3")
        table = doc.add_table(rows=conf_matrix.shape[0] + 1, cols=conf_matrix.shape[1] + 1)
        table.style = "Light Grid Accent 1"

        # Add header row
        header_cells = table.rows[0].cells
        header_cells[0].text = ""
        for i in range(conf_matrix.shape[1]):
            header_cells[i + 1].text = f"Predicted {i}"

        # Add data rows
        for i in range(conf_matrix.shape[0]):
            row_cells = table.rows[i + 1].cells
            row_cells[0].text = f"Actual {i}"
            for j in range(conf_matrix.shape[1]):
                row_cells[j + 1].text = str(conf_matrix[i, j])

    # Initialize the Word document
    doc = Document()
    doc.add_heading("XGBoost Survival Model Evaluation", level=1)

    # Initialize lists to store metrics
    auc_scores = []
    sens_all_years = []
    spec_all_years = []
    PPV = []
    NPV = []
    auc_low = []
    auc_upp = []

    # ROC Curve Plot
    plt.figure(figsize=(8, 6))
    plt.plot([0, 1], [0, 1], color='gray', linestyle='--', lw=1)  # Diagonal line (random)

    # Process results for each year
    for N in range(1, 11):
        days = 365 * N
        predicted_probs = 1 - np.array([sf_func(days, pred) for pred in bst.predict(dval)])

        # Create event observation column for the current year
        val[f"Event_obs_year{N}"] = [1 if time < days else 0 for time in val["Inc_Fx_Inc_fx_hip_days_to_cens"]]
        val.loc[val["Inc_Fx_Inc_fx_hip_bin_cens"] == 0, f"Event_obs_year{N}"] = 0

        y_test_event = val[f"Event_obs_year{N}"]

        # Calculate optimal cutoff
        optimal_cutoff = sensivity_specifity_cutoff(y_test_event, predicted_probs)
        doc.add_heading(f"Year {N} Results", level=2)
        doc.add_paragraph(f"Optimal Cut-off: {optimal_cutoff:.4f}")
    
        # Compute ROC curve and AUC
        fpr, tpr, _ = roc_curve(y_test_event, predicted_probs)
        roc_auc = auc(fpr, tpr)
        auc_scores.append(roc_auc)
        #ci auc
        low, upp = bootstrap_auc(predicted_probs,y_test_event, days ,n_bootstrap=200)
        auc_low.append(low)
        auc_upp.append(upp)
        doc.add_paragraph(f"AUC: {roc_auc:.4f} ({low:.4f}, {upp:.4f})")
        plt.plot(fpr, tpr, lw=2, label=f'Year {N} (AUC = {roc_auc:.2f})')

        # Apply cutoff to calculate binary predictions
        binary_predictions = (predicted_probs >= optimal_cutoff).astype(int)
        doc.add_paragraph(f"Calibration-in-the-large index: {np.mean(binary_predictions) - np.mean(y_test_event):.4f}")

        # Classification report
        report = classification_report(y_test_event, binary_predictions, target_names=["No Event", "Event"], output_dict=True)
        add_classification_report_to_doc(doc, report, N)

        # Confusion matrix
        conf_matrix = confusion_matrix(y_test_event, binary_predictions)
        add_confusion_matrix_to_doc(doc, conf_matrix, N)

        # Sensitivity and Specificity
        tp = conf_matrix[1, 1]
        fn = conf_matrix[1, 0]
        tn = conf_matrix[0, 0]
        fp = conf_matrix[0, 1]

        sensitivity = tp / (tp + fn) if (tp + fn) > 0 else 0
        specificity = tn / (tn + fp) if (tn + fp) > 0 else 0
        ppv_value   = tp / (tp + fp)
        npv_value   = tn / (tn + fn)
        

        df_met = bootstrap_metrics (binary_predictions, y_test_event)

        sens_all_years.append(sensitivity)
        spec_all_years.append(specificity)
        PPV.append ( ppv_value ) 
        NPV.append( npv_value )

        doc.add_paragraph(f"Sensitivity: {sensitivity:.4f} ({df_met.iloc[0,0]:.4f}, {df_met.iloc[0,1]:.4f})")
        doc.add_paragraph(f"Specificity: {specificity:.4f} ({df_met.iloc[0,2]:.4f}, {df_met.iloc[0,3]:.4f})") 
        doc.add_paragraph(f"Positive Predictive Value: {ppv_value:.4f} ({df_met.iloc[0,4]:.4f}, {df_met.iloc[0,5]:.4f})")
        doc.add_paragraph(f"Negative Predictive Value: {npv_value:.4f} ({df_met.iloc[0,6]:.4f}, {df_met.iloc[0,7]:.4f})")

        

    # Save ROC Curve
    plt.xlabel('False Positive Rate')
    plt.ylabel('True Positive Rate')
    plt.title('ROC Curve for All Years')
    plt.legend(loc="lower right")
    plt.grid(alpha=0.3)
    roc_buffer = io.BytesIO()
    plt.savefig(roc_buffer, format="png")
    plt.savefig(output_fig+"_ROC.pdf", format='pdf', dpi=1200)  # High-quality PDF format
    plt.savefig(output_fig+"_ROC.eps", format='eps', dpi=1200)

    plt.close()
    doc.add_heading("ROC Curve for All Years", level=1)
    doc.add_picture(roc_buffer, width=Inches(6))
    roc_buffer.close()

    # Plot AUC, Sensitivity, and Specificity
    plt.figure(figsize=(10, 6))
    plt.plot(range(1, 11), auc_scores, marker="o", label="AUC")
    plt.plot(range(1, 11), sens_all_years, marker="o", label="Sensitivity")
    plt.plot(range(1, 11), spec_all_years, marker="o", label="Specificity")
    plt.plot(range(1, 11), PPV, marker="o", label="PPV")
    plt.plot(range(1, 11), NPV, marker="o", label="NPV")
    
    
    plt.xlabel("Years")
    plt.ylabel("Metric Value")
    plt.title("AUC, Sensitivity, Specificity, PPV and NPV Over 10 Years")
    plt.legend()
    plt.grid(alpha=0.3)
    metrics_buffer = io.BytesIO()
    plt.savefig(metrics_buffer, format="png")
    plt.savefig(output_fig+"_AUC.pdf", format='pdf', dpi=1200)  # High-quality PDF format
    plt.savefig(output_fig+"_AUC.eps", format='eps', dpi=1200)

    plt.close()
    doc.add_heading("AUC, Sensitivity, Specificity, PPV and NPV Over 10 Years", level=1)
    doc.add_picture(metrics_buffer, width=Inches(6))
    metrics_buffer.close()
    #
    for N in range(1,11):
        days = 365*N
        predicted_probs = 1 - np.array([sf_func(days, pred) for pred in bst.predict(dval)])
    

        val[f"Event_obs_year{N}"] = [1 if time < days else 0 for time in val["Inc_Fx_Inc_fx_hip_days_to_cens"]]
        val.loc[val["Inc_Fx_Inc_fx_hip_bin_cens"] == 0, f"Event_obs_year{N}"] = 0

        y_test_event = val[f"Event_obs_year{N}"] 

        # Bin predicted probabilities
        n_bins = 10
        bins = np.linspace(0, 0.03, n_bins + 1)
        bin_indices = np.digitize(predicted_probs, bins) - 1
        # Calculate mean predicted and observed probabilities in each bin
        mean_predicted = []
        mean_observed = []

        for i in range(n_bins):
            # Get indices for current bin
            bin_mask = bin_indices == i
            if np.sum(bin_mask) > 0:
                # Mean predicted probability for the bin
                mean_predicted.append(np.mean(predicted_probs[bin_mask]))
                
                # Observed survival probability (Kaplan-Meier at the chosen time)
            
                observed_in_bin = np.mean(y_test_event[bin_mask])
                mean_observed.append(observed_in_bin)
            else:
                # Skip empty bins
                mean_predicted.append(np.nan)
                mean_observed.append(np.nan)

        # Drop NaNs for empty bins
        mean_predicted = np.array(mean_predicted)
        mean_observed = np.array(mean_observed)
        valid_mask = ~np.isnan(mean_predicted)
        plt.figure(figsize=(8, 6))
        plt.plot(mean_predicted[valid_mask], mean_observed[valid_mask],  marker='o')
        plt.plot([0, 0.3], [0, 0.3], 'k--')
        plt.xlabel("Predicted Survival Probability")
        plt.ylabel("Observed Survival Probability")
        plt.title(f"Calibration Plot at Year {N}")
        plt.legend()
        plt.grid(alpha=0.3)
        metrics_buffer = io.BytesIO()
        plt.savefig(metrics_buffer, format="png")
        plt.savefig(output_fig+"_Calibration_Year_"+str(N)+".pdf", format='pdf', dpi=1200)  # High-quality PDF format
        plt.savefig(output_fig+"_Calibration_Year_"+str(N)+".eps", format='eps', dpi=1200)

        plt.close()
        doc.add_heading(f"Calibration curve Year {N}", level=1)
        doc.add_picture(metrics_buffer, width=Inches(6))
        metrics_buffer.close()
    # Save the document
    doc.save(output_dir)
    print(f"Results saved to {output_dir}")


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Evaluate XGBoost model with validation dataset.")
    parser.add_argument("--input_dir", type=str, required=True, help="Path to the input data.")
    parser.add_argument("--model", type=str, required=True, help="Path of the model to evaluate.")
    parser.add_argument("--output_dir", type=str, required=True, help="Directory where the output will be saved.")
    parser.add_argument("--output_fig", type=str, required=True, help="Directory where Figures will be saved.")

    args = parser.parse_args()


    # Run the main function with the provided arguments
    main(args.input_dir, args.model, args.output_dir, args.output_fig)


